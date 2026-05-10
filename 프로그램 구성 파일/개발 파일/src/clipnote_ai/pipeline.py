from __future__ import annotations

import html
import json
import os
import re
import shutil
import tempfile
import time
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Callable
from urllib.parse import urlparse

from openai import OpenAI

from clipnote_ai.settings import AppSettings
from clipnote_ai.utils import (
    clamp_seconds,
    extract_json_object,
    find_ffmpeg,
    format_timecode,
    get_media_duration,
    parse_timecode,
    run_process,
    sanitize_filename,
    suggest_scene_count,
)


ProgressCallback = Callable[[str, float, str], None]
USER_PDF_NAME = "요약 노트.pdf"
USER_DOCX_NAME = "노션 붙여넣기용 요약 노트.docx"
USER_TRANSCRIPT_NAME = "전체 스크립트.txt"
SUPPORT_DIR_NAME = "기타 파일"
SUPPORT_MARKDOWN_NAME = "요약 노트.md"
SUPPORT_HTML_NAME = "요약 노트.html"
USD_TO_KRW = 1459.10
AUDIO_EXTENSIONS = {
    ".aac",
    ".aif",
    ".aiff",
    ".amr",
    ".caf",
    ".flac",
    ".m4a",
    ".m4b",
    ".mp3",
    ".oga",
    ".ogg",
    ".opus",
    ".wav",
    ".wma",
}
TRANSCRIPTION_CHUNK_SECONDS = 90
VIDEO_EXTENSIONS = {
    ".3g2",
    ".3gp",
    ".avi",
    ".flv",
    ".m2ts",
    ".m4v",
    ".mkv",
    ".mov",
    ".mp4",
    ".mpeg",
    ".mpg",
    ".mts",
    ".ogv",
    ".ts",
    ".webm",
    ".wmv",
}
MEDIA_EXTENSIONS = VIDEO_EXTENSIONS | AUDIO_EXTENSIONS

TEXT_MODEL_PRICING_USD_PER_1M = {
    "gpt-5.5": (5.00, 0.50, 30.00),
    "gpt-5.5-2026-04-23": (5.00, 0.50, 30.00),
    "gpt-5.4": (2.50, 0.25, 15.00),
    "gpt-5.4-2026-03-05": (2.50, 0.25, 15.00),
    "gpt-5.4-mini": (0.75, 0.075, 4.50),
    "gpt-5.4-mini-2026-03-17": (0.75, 0.075, 4.50),
    "gpt-5.4-nano": (0.20, 0.02, 1.25),
    "gpt-5.4-nano-2026-03-17": (0.20, 0.02, 1.25),
    "gpt-5": (1.25, 0.125, 10.00),
    "gpt-5-2025-08-07": (1.25, 0.125, 10.00),
    "gpt-5-mini": (0.25, 0.025, 2.00),
    "gpt-5-mini-2025-08-07": (0.25, 0.025, 2.00),
    "gpt-5-nano": (0.05, 0.005, 0.40),
    "gpt-5-nano-2025-08-07": (0.05, 0.005, 0.40),
    "gpt-4.1": (2.00, 0.50, 8.00),
    "gpt-4.1-2025-04-14": (2.00, 0.50, 8.00),
    "gpt-4.1-mini": (0.40, 0.10, 1.60),
    "gpt-4.1-mini-2025-04-14": (0.40, 0.10, 1.60),
    "gpt-4.1-nano": (0.10, 0.025, 0.40),
    "gpt-4.1-nano-2025-04-14": (0.10, 0.025, 0.40),
    "gpt-4o": (2.50, 1.25, 10.00),
    "gpt-4o-2024-11-20": (2.50, 1.25, 10.00),
    "gpt-4o-2024-08-06": (2.50, 1.25, 10.00),
    "gpt-4o-2024-05-13": (2.50, 1.25, 10.00),
    "gpt-4o-mini": (0.15, 0.075, 0.60),
    "gpt-4o-mini-2024-07-18": (0.15, 0.075, 0.60),
    "gpt-4-turbo": (10.00, 0.00, 30.00),
    "gpt-3.5-turbo": (0.50, 0.00, 1.50),
    "o4-mini": (1.10, 0.275, 4.40),
    "o4-mini-2025-04-16": (1.10, 0.275, 4.40),
    "o3": (2.00, 0.50, 8.00),
    "o3-2025-04-16": (2.00, 0.50, 8.00),
    "o3-mini": (1.10, 0.55, 4.40),
    "o3-mini-2025-01-31": (1.10, 0.55, 4.40),
    "o1": (15.00, 7.50, 60.00),
    "o1-2024-12-17": (15.00, 7.50, 60.00),
    "o1-mini": (1.10, 0.55, 4.40),
    "o1-mini-2024-09-12": (1.10, 0.55, 4.40),
    "o1-pro": (150.00, 0.00, 600.00),
}
TRANSCRIPTION_PRICING_USD_PER_MINUTE = {
    "gpt-4o-mini-transcribe": 0.003,
    "gpt-4o-transcribe": 0.006,
    "whisper-1": 0.006,
}


class UserFacingError(RuntimeError):
    """Error message that is already written for app users."""


@dataclass
class TranscriptChunk:
    index: int
    start: float
    end: float
    path: Path
    raw_text: str = ""
    clean_text: str = ""


@dataclass
class Scene:
    index: int
    seconds: int
    timecode: str
    heading: str
    summary: str
    quote: str
    why: str
    script: str = ""
    image_path: Path | None = None


@dataclass
class CostReport:
    text_input_tokens: int = 0
    text_cached_input_tokens: int = 0
    text_output_tokens: int = 0
    text_cost_usd: float = 0.0
    transcription_minutes: float = 0.0
    transcription_cost_usd: float = 0.0
    unknown_models: tuple[str, ...] = ()

    @property
    def total_cost_usd(self) -> float:
        return self.text_cost_usd + self.transcription_cost_usd

    @property
    def total_cost_krw(self) -> float:
        return self.total_cost_usd * USD_TO_KRW

    def format_for_log(self) -> str:
        krw = round(self.total_cost_krw)
        usd = self.total_cost_usd
        detail = (
            f"예상 API 비용: 약 {krw:,}원 "
            f"(${usd:.4f}, 환율 1달러={USD_TO_KRW:,.0f}원 기준)"
        )
        if self.unknown_models:
            detail += f"\n가격표에 없는 모델은 계산에서 제외됨: {', '.join(self.unknown_models)}"
        return detail


@dataclass
class PipelineResult:
    output_dir: Path
    video_path: Path
    transcript_path: Path
    summary_path: Path
    title: str
    cost_report: CostReport


class ApiCostTracker:
    def __init__(self) -> None:
        self.text_input_tokens = 0
        self.text_cached_input_tokens = 0
        self.text_output_tokens = 0
        self.text_cost_usd = 0.0
        self.transcription_minutes = 0.0
        self.transcription_cost_usd = 0.0
        self.unknown_models: set[str] = set()

    def add_transcription_minutes(self, model: str, minutes: float) -> None:
        rate = TRANSCRIPTION_PRICING_USD_PER_MINUTE.get(model)
        self.transcription_minutes += max(0.0, minutes)
        if rate is None:
            self.unknown_models.add(model)
            return
        self.transcription_cost_usd += max(0.0, minutes) * rate

    def add_text_usage(self, model: str, usage: object) -> None:
        input_tokens = self._usage_int(usage, ("input_tokens", "prompt_tokens"))
        output_tokens = self._usage_int(usage, ("output_tokens", "completion_tokens"))
        cached_tokens = self._cached_tokens(usage)

        self.text_input_tokens += input_tokens
        self.text_cached_input_tokens += cached_tokens
        self.text_output_tokens += output_tokens

        pricing = TEXT_MODEL_PRICING_USD_PER_1M.get(model)
        if pricing is None:
            self.unknown_models.add(model)
            return

        input_rate, cached_rate, output_rate = pricing
        billable_input = max(0, input_tokens - cached_tokens)
        self.text_cost_usd += (
            (billable_input * input_rate)
            + (cached_tokens * cached_rate)
            + (output_tokens * output_rate)
        ) / 1_000_000

    def report(self) -> CostReport:
        return CostReport(
            text_input_tokens=self.text_input_tokens,
            text_cached_input_tokens=self.text_cached_input_tokens,
            text_output_tokens=self.text_output_tokens,
            text_cost_usd=self.text_cost_usd,
            transcription_minutes=self.transcription_minutes,
            transcription_cost_usd=self.transcription_cost_usd,
            unknown_models=tuple(sorted(self.unknown_models)),
        )

    @staticmethod
    def _usage_int(usage: object, names: tuple[str, ...]) -> int:
        for name in names:
            value = getattr(usage, name, None)
            if value is None and isinstance(usage, dict):
                value = usage.get(name)
            if value is not None:
                try:
                    return int(value)
                except (TypeError, ValueError):
                    return 0
        return 0

    @staticmethod
    def _cached_tokens(usage: object) -> int:
        detail_names = ("input_tokens_details", "prompt_tokens_details")
        for detail_name in detail_names:
            details = getattr(usage, detail_name, None)
            if details is None and isinstance(usage, dict):
                details = usage.get(detail_name)
            if details is None:
                continue
            value = getattr(details, "cached_tokens", None)
            if value is None and isinstance(details, dict):
                value = details.get("cached_tokens")
            if value is not None:
                try:
                    return int(value)
                except (TypeError, ValueError):
                    return 0
        return 0


class VideoNotePipeline:
    def __init__(self, settings: AppSettings, progress: ProgressCallback | None = None):
        self.settings = settings
        self.progress = progress or (lambda _message, _percent, _detail: None)
        self.ffmpeg = find_ffmpeg()
        self.client = OpenAI(api_key=settings.api_key)
        self.costs = ApiCostTracker()

    @staticmethod
    def is_url(source: str) -> bool:
        parsed = urlparse(source.strip())
        return parsed.scheme in {"http", "https"} and bool(parsed.netloc)

    def run(self, source: str) -> PipelineResult:
        if not hasattr(self, "costs"):
            self.costs = ApiCostTracker()

        source = source.strip()
        if not source:
            raise ValueError("URL 또는 영상/오디오 파일을 입력해 주세요.")
        if not self.settings.api_key.strip():
            raise ValueError("OpenAI API 키를 입력해 주세요.")

        output_root = Path(self.settings.output_dir).expanduser().resolve()
        output_root.mkdir(parents=True, exist_ok=True)

        started = time.strftime("%y%m%d%H%M")

        self.progress("준비 중", 0.02, "저장할 파일 이름을 준비하고 있습니다.")

        source_kind = self._source_kind(source)
        if self.is_url(source):
            downloaded_path, source_title = self._download_video(source, started, output_root)
            source_video_path = downloaded_path
        else:
            source_video_path = Path(source).expanduser().resolve()
            if not source_video_path.exists():
                raise FileNotFoundError(f"영상 또는 오디오 파일을 찾을 수 없습니다: {source_video_path}")
            source_title = source_video_path.stem

        safe_title = sanitize_filename(source_title)
        final_base = self._unique_output_base(output_root, f"{started} {safe_title}", source_video_path.suffix or ".mp4")
        video_path = final_base.with_suffix(source_video_path.suffix or ".mp4")
        transcript_path = final_base.with_suffix(".txt")
        summary_path = final_base.with_name(f"{final_base.name}_요약").with_suffix(".txt")

        if source_video_path.resolve() != video_path.resolve():
            self.progress("파일 저장 중", 0.10, f"파일을 결과 폴더에 저장합니다: {video_path.name}")
            shutil.copy2(source_video_path, video_path)
            if self.is_url(source):
                try:
                    source_video_path.unlink()
                except OSError:
                    pass
        else:
            video_path = source_video_path

        duration = get_media_duration(video_path, self.ffmpeg)
        self.progress("파일 분석 중", 0.14, f"재생 길이: {format_timecode(duration)}")

        with tempfile.TemporaryDirectory(prefix="video_note_") as temp_dir:
            chunks = self._extract_audio_chunks(video_path, Path(temp_dir), duration)
            self._transcribe_chunks(chunks)
            self._clean_chunks(chunks)
            self._write_transcript(transcript_path, chunks)
            self._write_summary(summary_path, source_title, chunks, source_kind)

        self.progress("완료", 1.0, f"결과 생성 완료: {video_path.name}, {transcript_path.name}, {summary_path.name}")
        return PipelineResult(
            output_dir=output_root,
            video_path=video_path,
            transcript_path=transcript_path,
            summary_path=summary_path,
            title=source_title,
            cost_report=self.costs.report(),
        )

    def _unique_output_base(self, output_root: Path, base_name: str, video_suffix: str) -> Path:
        candidate = output_root / base_name
        suffix = 1
        while (
            candidate.with_suffix(video_suffix).exists()
            or candidate.with_suffix(".txt").exists()
            or candidate.with_name(f"{candidate.name}_요약").with_suffix(".txt").exists()
        ):
            suffix += 1
            candidate = output_root / f"{base_name} ({suffix})"
        return candidate

    def _source_kind(self, source: str) -> str:
        if not self.is_url(source):
            suffix = Path(source).suffix.lower()
            if suffix in AUDIO_EXTENSIONS:
                return "내 컴퓨터 오디오 파일"
            return "내 컴퓨터 영상 파일"

        host = urlparse(source.strip()).netloc.lower()
        if "instagram.com" in host:
            return "인스타그램 릴스 또는 짧은 세로 영상"
        if "youtube.com" in host or "youtu.be" in host:
            return "유튜브 영상"
        return "링크로 가져온 온라인 영상"

    def _best_source_title(self, info: dict[str, object]) -> str:
        raw_title = self._metadata_text(info, "title")
        candidates: list[str] = []
        if raw_title and not self._is_generic_online_title(raw_title):
            candidates.append(raw_title)

        for key in ("description", "caption", "alt_title", "fulltitle"):
            candidate = self._caption_title(self._metadata_text(info, key))
            if candidate and not self._is_generic_online_title(candidate):
                candidates.append(candidate)

        if raw_title:
            candidates.append(raw_title)

        for candidate in candidates:
            cleaned = self._caption_title(candidate)
            if cleaned:
                return cleaned
        return "downloaded_media"

    @staticmethod
    def _metadata_text(info: dict[str, object], key: str) -> str:
        value = info.get(key)
        if isinstance(value, str):
            return value.strip()
        return ""

    @staticmethod
    def _is_generic_online_title(title: str) -> bool:
        normalized = re.sub(r"\s+", " ", title.strip()).lower()
        return bool(
            re.fullmatch(r"(video|reel|post)( by [\w_.-]+)?", normalized)
            or re.fullmatch(r"instagram (video|reel|post)", normalized)
            or normalized in {"downloaded_video", "video", "reel", "post"}
        )

    @staticmethod
    def _caption_title(text: str, max_length: int = 80) -> str:
        if not text.strip():
            return ""

        text = re.sub(r"https?://\S+", "", text)
        text = re.sub(r"#\S+", "", text)
        text = re.sub(r"^\s*[\w.]+\s+on\s+Instagram:\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(
            r"^\s*[\d,]+\s+likes?,\s*[\d,]+\s+comments?\s*-\s*[^:]{1,80}:\s*",
            "",
            text,
            flags=re.IGNORECASE,
        )

        lines = [re.sub(r"\s+", " ", line).strip(" -|:") for line in text.splitlines()]
        title = next((line for line in lines if len(line) >= 2), "")
        if not title:
            return ""

        if len(title) <= max_length:
            return title

        boundary = max(title.rfind(mark, 0, max_length) for mark in (".", "?", "!", "。", "요", "다"))
        if boundary >= 20:
            return title[: boundary + 1].strip()
        return title[:max_length].rstrip() + "..."

    @staticmethod
    def _is_instagram_url(url: str) -> bool:
        return "instagram.com" in urlparse(url.strip()).netloc.lower()

    @staticmethod
    def _download_error_needs_cookies(error: BaseException) -> bool:
        message = str(error)
        lowered = message.lower()
        return any(
            phrase in lowered
            for phrase in (
                "login required",
                "cookies",
                "cookie",
                "rate-limit",
                "requested content is not available",
                "not available",
                "for authentication",
            )
        )

    def _cookie_browser_candidates(self) -> list[str]:
        selected = (self.settings.cookie_browser or "chrome").strip().lower()
        common = [
            selected,
            "chrome",
            "edge",
            "firefox",
            "brave",
            "chromium",
            "opera",
            "vivaldi",
            "whale",
        ]
        candidates: list[str] = []
        seen: set[str] = set()
        for browser in common:
            if browser and browser not in seen:
                candidates.append(browser)
                seen.add(browser)
        return candidates

    def _friendly_download_error(
        self,
        url: str,
        error: BaseException,
        retried_with_cookies: bool = False,
    ) -> UserFacingError:
        if self._is_instagram_url(url):
            browser = self.settings.cookie_browser or "chrome"
            if self.settings.use_browser_cookies or retried_with_cookies:
                return UserFacingError(
                    "Instagram에서 이 릴스를 바로 다운로드하지 못했습니다.\n\n"
                    f"'{browser}' 등 PC에 설치된 브라우저 쿠키로 다시 시도했지만, 로그인 정보가 없거나 Instagram이 요청을 막았습니다.\n\n"
                    "확인해 주세요.\n"
                    "1. Chrome 또는 Edge에서 Instagram에 로그인되어 있는지 확인\n"
                    "2. 비공개/삭제된 릴스가 아닌지 확인\n"
                    "3. 브라우저를 완전히 닫은 뒤 다시 실행\n"
                    "4. 계속 안 되면 릴스를 직접 저장한 영상 파일로 넣기"
                )
            return UserFacingError(
                "Instagram에서 로그인이 필요한 릴스라서 다운로드하지 못했습니다.\n\n"
                "해결 방법:\n"
                "1. Chrome 또는 Edge에서 Instagram에 로그인\n"
                "2. 프로그램의 '영상 가져오기'에서 '브라우저 쿠키 사용' 체크\n"
                "3. 로그인한 브라우저를 선택한 뒤 다시 실행\n\n"
                "그래도 안 되면 릴스를 직접 저장한 영상 파일로 넣어 주세요."
            )

        if retried_with_cookies:
            return UserFacingError(
                "링크 영상을 다운로드하지 못했습니다.\n\n"
                "PC에 설치된 브라우저 쿠키로 자동 재시도했지만 사이트에서 요청을 막았거나 로그인 정보가 부족합니다.\n\n"
                "확인해 주세요.\n"
                "1. 브라우저에서 해당 사이트에 로그인되어 있는지 확인\n"
                "2. 브라우저에서 영상이 정상 재생되는지 확인\n"
                "3. 브라우저를 완전히 닫은 뒤 다시 실행\n"
                "4. 계속 안 되면 영상을 직접 저장한 파일로 넣기"
            )

        return UserFacingError(
            "링크 영상을 다운로드하지 못했습니다.\n\n"
            "가능한 원인:\n"
            "- 영상이 비공개이거나 삭제됨\n"
            "- 사이트에서 로그인 또는 쿠키를 요구함\n"
            "- 짧은 시간에 요청이 많아 일시적으로 막힘\n\n"
            "브라우저에서 영상이 정상 재생되는지 확인한 뒤 다시 시도해 주세요."
        )

    def _download_with_ytdlp(
        self,
        yt_dlp_module: object,
        url: str,
        ydl_opts: dict[str, object],
        downloads_dir: Path,
        started: str,
    ) -> tuple[Path, str]:
        with yt_dlp_module.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            if info is None:
                raise RuntimeError("영상 정보를 가져오지 못했습니다.")
            title = self._best_source_title(info)
            downloaded = Path(ydl.prepare_filename(info))
            merged = downloaded.with_suffix(".mp4")
            if merged.exists():
                downloaded = merged
            if not downloaded.exists():
                candidates = sorted(downloads_dir.glob(f"{started}_*"), key=lambda path: path.stat().st_mtime, reverse=True)
                if not candidates:
                    raise RuntimeError("다운로드된 영상 파일을 찾지 못했습니다.")
                downloaded = candidates[0]
        return downloaded.resolve(), title

    def _download_with_browser_cookie_fallback(
        self,
        yt_dlp_module: object,
        url: str,
        base_opts: dict[str, object],
        downloads_dir: Path,
        started: str,
        tried_browsers: set[str] | None = None,
    ) -> tuple[Path, str]:
        tried_browsers = tried_browsers or set()
        last_error: Exception | None = None
        for browser in self._cookie_browser_candidates():
            if browser in tried_browsers:
                continue
            self.progress(
                "브라우저 쿠키 자동 확인 중",
                0.07,
                f"{browser}에 저장된 로그인 정보로 링크 영상을 다시 시도합니다.",
            )
            retry_opts = dict(base_opts)
            retry_opts["cookiesfrombrowser"] = (browser,)
            try:
                return self._download_with_ytdlp(yt_dlp_module, url, retry_opts, downloads_dir, started)
            except Exception as exc:
                last_error = exc
                self.progress(
                    "다른 브라우저 쿠키 확인 중",
                    0.07,
                    f"{browser} 쿠키로는 다운로드하지 못했습니다. 다른 브라우저를 확인합니다.",
                )
        if last_error is not None:
            raise last_error
        raise RuntimeError("사용 가능한 브라우저 쿠키 후보가 없습니다.")

    def _download_video(self, url: str, started: str, downloads_dir: Path) -> tuple[Path, str]:
        downloads_dir.mkdir(parents=True, exist_ok=True)
        self.progress(
            "영상 다운로드 중",
            0.05,
            f"링크 영상을 결과 폴더에 저장합니다: {downloads_dir}",
        )
        import yt_dlp

        outtmpl = str(downloads_dir / f"__download_{started}_%(title).90s.%(ext)s")
        ydl_opts: dict[str, object] = {
            "format": "best[ext=mp4][vcodec!=none][acodec!=none]/best[vcodec!=none][acodec!=none]/bv*+ba/b",
            "outtmpl": outtmpl,
            "merge_output_format": "mp4",
            "noplaylist": True,
            "quiet": True,
            "no_warnings": True,
            "windowsfilenames": True,
            "ffmpeg_location": str(self.ffmpeg),
        }
        tried_browsers: set[str] = set()
        if self.settings.use_browser_cookies:
            browser = (self.settings.cookie_browser or "chrome").strip().lower()
            ydl_opts["cookiesfrombrowser"] = (browser,)
            tried_browsers.add(browser)

        try:
            downloaded, title = self._download_with_ytdlp(yt_dlp, url, ydl_opts, downloads_dir, started)
        except UserFacingError:
            raise
        except Exception as exc:
            should_retry_with_cookies = self._download_error_needs_cookies(exc)
            if not should_retry_with_cookies:
                raise self._friendly_download_error(url, exc) from exc

            self.progress(
                "브라우저 쿠키 자동 재시도 중",
                0.07,
                "로그인이 필요할 수 있어 PC 브라우저에 저장된 쿠키를 자동으로 확인합니다.",
            )
            try:
                downloaded, title = self._download_with_browser_cookie_fallback(
                    yt_dlp,
                    url,
                    ydl_opts,
                    downloads_dir,
                    started,
                    tried_browsers,
                )
            except Exception as retry_exc:
                raise self._friendly_download_error(url, retry_exc, retried_with_cookies=True) from retry_exc
        self.progress("영상 다운로드 완료", 0.09, f"저장된 영상: {downloaded}")
        return downloaded.resolve(), title

    def _extract_audio_chunks(self, video_path: Path, support_dir: Path, duration: float) -> list[TranscriptChunk]:
        self.progress("음성 추출 중", 0.18, "오디오를 전사용 작은 조각으로 나누고 있습니다.")
        chunk_dir = support_dir / "audio_chunks"
        chunk_dir.mkdir(parents=True, exist_ok=True)
        chunk_seconds = TRANSCRIPTION_CHUNK_SECONDS
        output_pattern = chunk_dir / "chunk_%03d.mp3"
        completed = run_process(
            [
                self.ffmpeg,
                "-y",
                "-i",
                video_path,
                "-map",
                "0:a:0",
                "-vn",
                "-ac",
                "1",
                "-ar",
                "16000",
                "-b:a",
                "48k",
                "-f",
                "segment",
                "-segment_time",
                str(chunk_seconds),
                "-reset_timestamps",
                "1",
                output_pattern,
            ]
        )
        if completed.returncode != 0:
            raise RuntimeError(f"오디오 추출에 실패했습니다.\n{completed.stderr[-1200:]}")

        chunk_paths = sorted(chunk_dir.glob("chunk_*.mp3"))
        if not chunk_paths:
            raise RuntimeError("추출된 오디오 조각이 없습니다. 영상/오디오 파일에 음성이 있는지 확인해 주세요.")

        chunks: list[TranscriptChunk] = []
        elapsed = 0.0
        for index, path in enumerate(chunk_paths):
            start = elapsed
            try:
                chunk_duration = get_media_duration(path, self.ffmpeg)
            except Exception:
                chunk_duration = min(chunk_seconds, max(0.0, duration - start))
            end = min(duration, start + max(0.0, chunk_duration))
            if end <= start:
                end = min(duration, start + chunk_seconds)
            chunks.append(TranscriptChunk(index=index, start=start, end=end, path=path))
            elapsed = end
        return chunks

    def _transcribe_chunks(self, chunks: list[TranscriptChunk]) -> None:
        previous_tail = ""
        total = len(chunks)
        for chunk in chunks:
            base_percent = 0.24 + (chunk.index / max(1, total)) * 0.26
            self.progress(
                "전사 중",
                base_percent,
                f"{chunk.index + 1}/{total} 조각 전사: {format_timecode(chunk.start)}-{format_timecode(chunk.end)}",
            )
            chunk.raw_text = self._transcribe_file(chunk.path, previous_tail).strip()
            self.costs.add_transcription_minutes(
                self.settings.transcription_model,
                max(0.0, chunk.end - chunk.start) / 60,
            )
            previous_tail = chunk.raw_text[-600:]

    def _transcribe_file(self, path: Path, previous_tail: str) -> str:
        prompt = (
            "한국어 중심 영상일 수 있습니다. 자연스러운 띄어쓰기와 문장부호를 살려 주세요. "
            "브랜드명, 사이트명, 앱 이름, 버튼명, 기능명, 코드, URL, 전문 약어처럼 실제 표기가 중요한 말은 "
            "알파벳 표기를 보존해 주세요. 다만 일반 명사나 한국어 대화에서 보통 한글로 쓰는 말은 "
            "억지로 영어로 바꾸지 말고 자연스러운 한글 표기를 우선해 주세요."
        )
        if previous_tail:
            prompt += f"\n직전 내용 일부: {previous_tail}"

        attempts = [
            {"prompt": prompt},
            {"language": "ko", "prompt": prompt},
            {},
        ]
        last_error: Exception | None = None
        for extra in attempts:
            try:
                with path.open("rb") as audio_file:
                    result = self.client.audio.transcriptions.create(
                        model=self.settings.transcription_model,
                        file=audio_file,
                        response_format="text",
                        **extra,
                    )
                if isinstance(result, str):
                    return result
                return str(getattr(result, "text", result))
            except Exception as exc:
                last_error = exc
        raise RuntimeError(f"OpenAI 전사 호출에 실패했습니다: {last_error}") from last_error

    def _clean_chunks(self, chunks: list[TranscriptChunk]) -> None:
        settings = getattr(self, "settings", None)
        if not bool(getattr(settings, "polish_transcript", True)):
            self.progress(
                "전사문 정리 건너뜀",
                0.68,
                "비용 절약을 위해 맞춤법/띄어쓰기 정리를 건너뛰고 전사 원문을 보존합니다.",
            )
            for chunk in chunks:
                chunk.clean_text = chunk.raw_text
            return

        total = len(chunks)
        for chunk in chunks:
            base_percent = 0.50 + (chunk.index / max(1, total)) * 0.18
            self.progress(
                "문장 다듬는 중",
                base_percent,
                f"{chunk.index + 1}/{total} 조각의 맞춤법과 띄어쓰기를 정리하고 있습니다.",
            )
            if not chunk.raw_text.strip():
                chunk.clean_text = ""
                continue
            chunk.clean_text = self._clean_chunk_text(chunk).strip()

    def _clean_chunk_text(self, chunk: TranscriptChunk) -> str:
        system = (
            "너는 한국어 영상 전사문 교정자다. 의미를 바꾸거나 내용을 추가하지 말고, "
            "맞춤법, 띄어쓰기, 문장부호, 어색한 ASR 오류만 자연스럽게 고친다. "
            "절대 요약하지 말고, 입력된 모든 문장을 빠짐없이 같은 순서로 보존한다. "
            "브랜드명, 사이트명, 앱 이름, 버튼명, 기능명, 코드, URL, 전문 약어처럼 실제 표기가 중요한 말만 "
            "알파벳 표기를 보존한다. 일반 명사, 품종명, 직업명, 감정 표현처럼 한국어 문장 안에서 "
            "보통 한글로 쓰는 말은 영어로 바꾸지 말고 자연스러운 한글 표기를 우선한다. "
            "예: Doberman은 특별히 영문 표기를 말하는 맥락이 아니면 도베르만으로 쓴다. "
            "ASR이 브랜드명이나 제품명을 한글로 적었고 문맥상 영어 표기가 명확한 경우에만 영어로 복원한다. "
            "확실하지 않은 고유명사나 일반 단어는 추측해서 새 이름을 만들지 않는다. "
            "결과 텍스트만 반환한다."
        )
        user = f"구간: {format_timecode(chunk.start)}-{format_timecode(chunk.end)}\n\n{chunk.raw_text}"
        cleaned = self._text_response(system=system, user=user).strip()
        if not self._is_suspiciously_short_clean_text(chunk.raw_text, cleaned):
            return cleaned

        retry = self._text_response(
            system=(
                system
                + " 이전 응답이 원문보다 너무 짧으면 실패다. 이번에는 한 문장도 생략하지 말고 전체 전사문을 모두 반환한다."
            ),
            user=user,
        ).strip()
        if not self._is_suspiciously_short_clean_text(chunk.raw_text, retry):
            return retry

        self.progress(
            "문장 다듬기 보정",
            0.68,
            f"{chunk.index + 1}번 조각의 정리 결과가 너무 짧아 원문 전사문을 보존합니다.",
        )
        return chunk.raw_text

    def _is_suspiciously_short_clean_text(self, raw_text: str, clean_text: str) -> bool:
        raw = self._strip_transcript_labels(raw_text)
        clean = self._strip_transcript_labels(clean_text)
        raw_compact = re.sub(r"\s+", "", raw)
        clean_compact = re.sub(r"\s+", "", clean)
        if len(raw_compact) < 260:
            return False
        if not clean_compact:
            return True
        if len(clean_compact) < len(raw_compact) * 0.60:
            return True

        raw_sentence_count = len(self._split_sentences(raw))
        clean_sentence_count = len(self._split_sentences(clean))
        return raw_sentence_count >= 6 and clean_sentence_count < raw_sentence_count * 0.50

    def _text_response(self, system: str, user: str) -> str:
        try:
            response = self.client.responses.create(
                model=self.settings.text_model,
                input=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
            )
            self.costs.add_text_usage(self.settings.text_model, getattr(response, "usage", None))
            output_text = getattr(response, "output_text", None)
            if output_text:
                return str(output_text)
            pieces: list[str] = []
            for item in getattr(response, "output", []) or []:
                for content in getattr(item, "content", []) or []:
                    text = getattr(content, "text", None)
                    if text:
                        pieces.append(str(text))
            if pieces:
                return "\n".join(pieces)
        except Exception:
            pass

        completion = self.client.chat.completions.create(
            model=self.settings.text_model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )
        self.costs.add_text_usage(self.settings.text_model, getattr(completion, "usage", None))
        return completion.choices[0].message.content or ""

    def _write_transcript(self, transcript_path: Path, chunks: list[TranscriptChunk]) -> Path:
        paragraphs: list[str] = []
        for chunk in chunks:
            text = self._strip_transcript_labels(chunk.clean_text or chunk.raw_text)
            paragraphs.extend(self._note_paragraphs(text))
        transcript_path.write_text("\n\n".join(paragraphs).strip() + "\n", encoding="utf-8")
        return transcript_path

    def _write_summary(self, summary_path: Path, title: str, chunks: list[TranscriptChunk], source_kind: str = "동영상") -> Path:
        self.progress("요약 정리 중", 0.86, "전체 스크립트를 기준으로 적절한 길이의 상세 요약을 만들고 있습니다.")
        transcript = self._summary_source_text(chunks)
        target_sentences = self._summary_target_sentence_count(transcript)
        summary_mode = self._summary_mode_instruction(transcript, source_kind)
        summary = self._text_response(
            system=(
                "너는 한국어 영상 스크립트를 정리하는 전문 편집자다. "
                "원문을 그대로 베끼지 않고, 핵심 의미와 행동을 읽기 쉬운 요약문으로 재구성한다. "
                "원문에 없는 내용, 추측, 평가를 추가하지 않는다. "
                "영상의 성격에 맞춰 정보형 영상은 개념, 절차, 근거, 수치, 조건, 결론을 중심으로 정리하고, "
                "브이로그, 릴스, 홍보, 대화형 영상은 사건 흐름, 맥락, 핵심 장면, 주장, 분위기를 중심으로 정리한다."
            ),
            user=(
                f"영상 제목: {title}\n\n"
                f"영상 유형: {source_kind}\n"
                f"목표 길이: 약 {target_sentences}문장\n\n"
                f"{summary_mode}\n\n"
                "아래 전사문 전체를 입력으로 삼아 요약해 주세요.\n\n"
                "요약 원칙:\n"
                "- 핵심 수치, 금액, 날짜, 기간, 비율, 조건, 인물/회사/제품명, 단계, 예외, 원인과 결과는 반드시 남긴다.\n"
                "- 반복 표현, 말버릇, 중복 설명, 진행자가 시간을 끄는 말만 줄인다.\n"
                "- 단순히 비례해서 줄이지 말고, 중요한 정보 밀도가 높은 부분은 길게 남긴다.\n"
                "- 원문 문장을 그대로 나열하지 말고, 사용자가 바로 이해할 수 있게 자연스럽게 압축한다.\n"
                "- 목표 길이를 크게 벗어나지 않되, 중요한 디테일을 버려야 할 정도로 억지로 줄이지는 않는다.\n"
                "- 원문의 순서를 최대한 유지한다.\n"
                "- 확실하지 않은 내용은 단정하지 않는다.\n\n"
                "출력 형식:\n"
                "1. 제목, 머리말, '요약' 같은 라벨 없이 바로 본문으로 시작한다.\n"
                "2. 한 문장 또는 짧은 의미 단위마다 줄을 나눠 읽기 쉽게 쓴다.\n"
                "3. 너무 긴 문단은 만들지 않는다. 문단 사이에는 빈 줄을 하나 넣는다.\n"
                "4. 불릿과 번호 목록을 남발하지 말고, 필요한 경우에만 짧게 사용한다.\n\n"
                f"전사문:\n{transcript}"
            ),
        ).strip()
        summary_path.write_text(self._normalize_summary_text(summary, title), encoding="utf-8")
        return summary_path

    def _summary_source_text(self, chunks: list[TranscriptChunk], max_chars: int = 180_000) -> str:
        blocks: list[str] = []
        for chunk in chunks:
            text = self._strip_transcript_labels(chunk.clean_text or chunk.raw_text)
            blocks.append(text)
        transcript = "\n\n".join(block for block in blocks if block.strip()).strip()
        if len(transcript) <= max_chars:
            return transcript

        head = transcript[: max_chars // 2].rstrip()
        tail = transcript[-max_chars // 2 :].lstrip()
        return f"{head}\n\n...[전체 전사문이 길어 중간 일부를 줄였습니다]...\n\n{tail}"

    def _summary_target_sentence_count(self, transcript: str) -> int:
        if not transcript.strip():
            return 0

        if not bool(getattr(self.settings, "auto_summary_sentences", True)):
            return self._clamp_int(getattr(self.settings, "summary_sentence_count", 30), 3, 160)

        source_sentence_count = len(self._split_sentences(transcript))
        if source_sentence_count <= 0:
            rough_count = max(1, len(transcript) // 70)
            source_sentence_count = rough_count

        if source_sentence_count <= 4:
            return 1 if len(transcript) <= 260 else 2
        if source_sentence_count <= 8:
            return 2
        if source_sentence_count <= 12:
            return 3
        return self._clamp_int(round(source_sentence_count / 5), 6, 120)

    def _summary_mode_instruction(self, transcript: str, source_kind: str) -> str:
        sentence_count = len(self._split_sentences(transcript))
        is_short = sentence_count <= 8 or len(transcript) <= 520 or "릴스" in source_kind
        if is_short:
            return (
                "짧은 영상 요약 방식: 원문을 문장별로 다시 쓰지 말고, "
                "영상이 알려주는 핵심 행동/절차/결론을 1-2문장으로 압축한다. "
                "튜토리얼이면 '무엇을 하려면 어떤 앱/버튼/단계를 거치면 된다' 형태로 정리한다."
            )
        return (
            "긴 영상 요약 방식: 전체 흐름을 보존하되 반복과 잡담을 줄이고, "
            "중요한 주장, 근거, 단계, 수치, 예외를 중심으로 상세하게 정리한다."
        )

    @staticmethod
    def _clamp_int(value: object, low: int, high: int) -> int:
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            parsed = low
        return max(low, min(high, parsed))

    def _normalize_summary_text(self, summary: str, title: str) -> str:
        cleaned = re.sub(r"\n{3,}", "\n\n", summary.strip())
        if not cleaned:
            return "\n"

        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
        lines = self._drop_summary_heading(lines, title)
        body = " ".join(lines).strip()
        if not body:
            return "\n"

        paragraphs = self._note_paragraphs(body)
        return "\n\n".join(paragraphs).strip() + "\n"

    def _drop_summary_heading(self, lines: list[str], title: str) -> list[str]:
        if not lines:
            return lines

        first = lines[0].strip()
        if re.fullmatch(r"(요약|summary)", first, flags=re.IGNORECASE):
            return lines[1:]

        if re.match(r"^(영상\s*제목|제목)\s*[:：]", first):
            return lines[1:]

        summary_label = re.match(r"^요약\s*[:：]\s*(.+)$", first)
        if summary_label:
            rest = summary_label.group(1).strip()
            return ([rest] if rest else []) + lines[1:]

        first_key = self._compact_heading(first)
        title_key = self._compact_heading(title)
        if first_key and title_key:
            is_same_title = (
                first_key == title_key
                or first_key in title_key
                or title_key in first_key
                or SequenceMatcher(None, first_key, title_key).ratio() >= 0.78
            )
            if is_same_title:
                return lines[1:]

        return lines

    @staticmethod
    def _compact_heading(text: str) -> str:
        return re.sub(r"[\W_]+", "", text, flags=re.UNICODE).lower()

    def _strip_transcript_labels(self, text: str) -> str:
        cleaned = re.sub(r"^\s*\[[0-9:.]+\s*-\s*[0-9:.]+\]\s*", "", text.strip())
        cleaned = re.sub(r"^\s*구간:\s*[0-9:.]+\s*-\s*[0-9:.]+\s*", "", cleaned)
        return cleaned.strip()

    def _split_sentences(self, text: str) -> list[str]:
        normalized = re.sub(r"\s+", " ", text.strip())
        normalized = re.sub(r"(?<=\d)\.\s+(?=\d)", ".", normalized)
        if not normalized:
            return []

        sentences: list[str] = []
        start = 0
        index = 0
        while index < len(normalized):
            char = normalized[index]
            if char not in ".!?。！？":
                index += 1
                continue

            if char == "." and self._is_non_sentence_period(normalized, index, start):
                index += 1
                continue

            end = index + 1
            while end < len(normalized) and normalized[end] in "\"')]}”’》」』":
                end += 1

            sentence = normalized[start:end].strip()
            if sentence:
                sentences.append(sentence)

            start = end
            while start < len(normalized) and normalized[start].isspace():
                start += 1
            index = start

        tail = normalized[start:].strip()
        if tail:
            sentences.append(tail)

        if len(sentences) == 1 and len(sentences[0]) > 360:
            return self._split_long_text(sentences[0], 220)
        return sentences

    def _is_non_sentence_period(self, text: str, index: int, sentence_start: int) -> bool:
        previous = text[index - 1] if index > 0 else ""
        next_immediate = text[index + 1] if index + 1 < len(text) else ""
        next_index = index + 1
        while next_index < len(text) and text[next_index].isspace():
            next_index += 1
        next_nonspace = text[next_index] if next_index < len(text) else ""

        if previous.isdigit() and next_nonspace.isdigit():
            return True

        token_before_period = text[sentence_start:index].strip()
        if re.fullmatch(r"\d{1,3}", token_before_period) and next_nonspace:
            return True

        if (
            previous.isascii()
            and previous.isalnum()
            and next_immediate.isascii()
            and next_immediate.isalnum()
        ):
            return True

        return False

    def _transcript_paragraphs(self, text: str, sentences_per_paragraph: int = 2, max_chars: int = 260) -> list[str]:
        sentences = self._split_sentences(text)
        if not sentences:
            return []

        paragraphs: list[str] = []
        current: list[str] = []
        current_len = 0

        for sentence in sentences:
            current.append(sentence)
            current_len += len(sentence)
            if len(current) >= sentences_per_paragraph or current_len >= max_chars:
                paragraphs.append(" ".join(current).strip())
                current = []
                current_len = 0

        if current:
            paragraphs.append(" ".join(current).strip())

        if len(paragraphs) == 1 and len(paragraphs[0]) > max_chars:
            paragraphs = self._split_long_text(paragraphs[0], max_chars)

        return paragraphs

    def _note_paragraphs(self, text: str) -> list[str]:
        sentences = self._split_sentences(text)
        if not sentences:
            return []

        paragraphs: list[str] = []
        for sentence in sentences:
            if len(sentence) > 210:
                paragraphs.extend(self._split_long_text(sentence, 210))
            else:
                paragraphs.append(sentence)
        return paragraphs

    def _split_long_text(self, text: str, max_chars: int = 260) -> list[str]:
        words = text.split(" ")
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if len(candidate) > max_chars and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        return lines or [text]

    def _transcript_for_prompt(self, chunks: list[TranscriptChunk], max_chars: int = 170_000) -> str:
        per_chunk_limit = max(900, max_chars // max(1, len(chunks)))
        blocks: list[str] = []
        for chunk in chunks:
            text = (chunk.clean_text or chunk.raw_text).strip()
            if len(text) > per_chunk_limit:
                text = text[:per_chunk_limit].rstrip() + "\n...[긴 구간이라 일부 생략됨]"
            blocks.append(f"[{format_timecode(chunk.start)} - {format_timecode(chunk.end)}]\n{text}")
        return "\n\n".join(blocks)

    def _analyze_scenes(
        self,
        source_title: str,
        source_label: str,
        duration: float,
        chunks: list[TranscriptChunk],
    ) -> dict[str, object]:
        auto_count = suggest_scene_count(
            duration,
            int(self.settings.min_scene_count),
            int(self.settings.max_scene_count),
        )
        scene_instruction = (
            f"자동 추천 장면 수는 {auto_count}개다. 내용 밀도에 따라 "
            f"{self.settings.min_scene_count}-{self.settings.max_scene_count}개 사이에서 정하라."
            if self.settings.auto_scene_count
            else f"반드시 {self.settings.fixed_scene_count}개 장면을 고르라."
        )
        transcript = self._transcript_for_prompt(chunks)
        self.progress("주요 장면 고르는 중", 0.72, "AI가 타임라인에서 중요한 순간을 고르고 있습니다.")
        json_text = self._text_response(
            system=(
                "너는 한국어 영상 편집자이자 노트 작성자다. 전사문을 읽고 문맥상 중요한 장면을 고른다. "
                "반드시 JSON만 반환한다. 설명 문장이나 마크다운 코드는 쓰지 않는다."
            ),
            user=(
                f"원본 제목: {source_title}\n"
                f"원본: {source_label}\n"
                f"영상 길이: {format_timecode(duration)} ({int(duration)}초)\n"
                f"{scene_instruction}\n\n"
                "장면 선택 기준:\n"
                "- 핵심 주장, 전환점, 예시, 결론, 감정적으로 강한 순간을 우선한다.\n"
                "- 릴스처럼 짧으면 장면 간격을 촘촘히, 강의처럼 길면 챕터처럼 넓게 배치한다.\n"
                "- 같은 말을 반복하는 장면은 하나만 고른다.\n"
                "- seconds는 0 이상 영상 길이 미만의 정수여야 한다.\n\n"
                "장면 아래에 넣을 대본(script) 기준:\n"
                "- 선택한 장면과 직접 관련된 전사문 발췌문만 넣는다.\n"
                "- 요약하거나 새 문장을 만들지 말고, 교정된 전사문 표현을 최대한 그대로 사용한다.\n"
                "- 보통 2~8문장 정도로, 블로그 글에서 이미지 아래에 붙일 본문처럼 자연스럽게 고른다.\n"
                "- 설명, 선정 이유, 대표 발화 같은 별도 코멘트는 넣지 않는다.\n\n"
                "JSON 형식:\n"
                "{\n"
                '  "title": "노트 제목",\n'
                '  "scenes": [\n'
                '    {"seconds": 12, "timecode": "00:00:12", "heading": "장면 제목", "script": "이미지 아래에 넣을 전사문 발췌"}\n'
                "  ]\n"
                "}\n\n"
                f"전사문:\n{transcript}"
            ),
        )
        return extract_json_object(json_text)

    def _normalize_scenes(
        self,
        analysis: dict[str, object],
        duration: float,
        chunks: list[TranscriptChunk],
    ) -> list[Scene]:
        raw_scenes = analysis.get("scenes")
        scenes: list[Scene] = []
        if isinstance(raw_scenes, list):
            for item in raw_scenes:
                if not isinstance(item, dict):
                    continue
                seconds = item.get("seconds")
                if seconds is None and isinstance(item.get("timecode"), str):
                    seconds = parse_timecode(str(item["timecode"]))
                if seconds is None:
                    continue
                safe_seconds = clamp_seconds(float(seconds), duration)
                scenes.append(
                    Scene(
                        index=len(scenes) + 1,
                        seconds=safe_seconds,
                        timecode=format_timecode(safe_seconds),
                        heading=str(item.get("heading") or f"주요 장면 {len(scenes) + 1}"),
                        summary="",
                        quote="",
                        why="",
                        script=str(item.get("script") or item.get("transcript") or ""),
                    )
                )

        if not scenes:
            count = (
                int(self.settings.fixed_scene_count)
                if not self.settings.auto_scene_count
                else suggest_scene_count(duration, self.settings.min_scene_count, self.settings.max_scene_count)
            )
            step = duration / (count + 1)
            for index in range(count):
                seconds = clamp_seconds((index + 1) * step, duration)
                scenes.append(
                    Scene(
                        index=index + 1,
                        seconds=seconds,
                        timecode=format_timecode(seconds),
                        heading=f"주요 장면 {index + 1}",
                        summary="",
                        quote="",
                        why="",
                        script="",
                    )
                )

        scenes = sorted(scenes, key=lambda scene: scene.seconds)
        for index, scene in enumerate(scenes, start=1):
            scene.index = index
            scene.timecode = format_timecode(scene.seconds)
            if not scene.script.strip():
                scene.script = self._fallback_scene_script(scene.seconds, chunks)
        return scenes

    def _attach_full_transcript_to_scenes(self, scenes: list[Scene], chunks: list[TranscriptChunk]) -> None:
        if not scenes:
            return

        ordered = sorted(scenes, key=lambda scene: scene.seconds)
        assigned: dict[int, list[str]] = {scene.index: [] for scene in ordered}

        def scene_for_time(seconds: float) -> Scene:
            selected = ordered[0]
            for scene in ordered:
                if seconds >= scene.seconds:
                    selected = scene
                else:
                    break
            return selected

        for chunk in chunks:
            text = (chunk.clean_text or chunk.raw_text).strip()
            sentences = self._split_sentences(text)
            if not sentences:
                continue
            duration = max(1.0, float(chunk.end) - float(chunk.start))
            for sentence_index, sentence in enumerate(sentences):
                sentence_time = float(chunk.start) + duration * ((sentence_index + 0.5) / len(sentences))
                assigned[scene_for_time(sentence_time).index].append(sentence)

        for scene in ordered:
            full_script = " ".join(assigned.get(scene.index, [])).strip()
            if full_script:
                scene.script = full_script

    def _fallback_scene_script(self, seconds: float, chunks: list[TranscriptChunk]) -> str:
        if not chunks:
            return ""
        chunk = next((item for item in chunks if item.start <= seconds <= item.end), None)
        if chunk is None:
            chunk = min(chunks, key=lambda item: abs(((item.start + item.end) / 2) - seconds))
        paragraphs = self._transcript_paragraphs(chunk.clean_text or chunk.raw_text)
        return "\n\n".join(paragraphs[:4]).strip()

    def _extract_scene_images(self, video_path: Path, support_dir: Path, scenes: list[Scene]) -> None:
        frames_dir = support_dir / "frames"
        frames_dir.mkdir(parents=True, exist_ok=True)
        total = len(scenes)
        for scene in scenes:
            self.progress(
                "이미지 추출 중",
                0.78 + (scene.index / max(1, total)) * 0.12,
                f"{scene.index}/{total} 장면 이미지 저장: {scene.timecode}",
            )
            output = frames_dir / f"scene_{scene.index:02d}_{scene.timecode.replace(':', '-')}.jpg"
            completed = run_process(
                [
                    self.ffmpeg,
                    "-y",
                    "-ss",
                    f"{scene.seconds:.3f}",
                    "-i",
                    video_path,
                    "-frames:v",
                    "1",
                    "-q:v",
                    "2",
                    output,
                ]
            )
            if completed.returncode != 0 or not output.exists():
                raise RuntimeError(f"장면 이미지 추출에 실패했습니다.\n{completed.stderr[-1200:]}")
            scene.image_path = output

    def _source_link(self, source_label: str) -> str:
        return source_label if self.is_url(source_label) else ""

    def _render_markdown(
        self,
        job_dir: Path,
        source_title: str,
        source_label: str,
        duration: float,
        chunks: list[TranscriptChunk],
        scenes: list[Scene],
        analysis: dict[str, object],
    ) -> Path:
        markdown_path = job_dir / SUPPORT_MARKDOWN_NAME
        title = str(analysis.get("title") or source_title)
        source_link = self._source_link(source_label)
        lines = [f"# {title}", ""]
        if source_link:
            lines.extend([source_link, ""])

        for scene in scenes:
            rel_image = scene.image_path.relative_to(markdown_path.parent).as_posix() if scene.image_path else ""
            lines.extend(
                [
                    f"## {scene.heading} ({scene.timecode})",
                    "",
                    f"![{scene.heading}]({rel_image})" if rel_image else "",
                    "",
                ]
            )
            lines.extend(self._note_paragraphs(scene.script))
            lines.append("")

        markdown_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        return markdown_path

    def _render_html(
        self,
        job_dir: Path,
        source_title: str,
        source_label: str,
        duration: float,
        chunks: list[TranscriptChunk],
        scenes: list[Scene],
        analysis: dict[str, object],
    ) -> Path:
        html_path = job_dir / SUPPORT_HTML_NAME
        title = str(analysis.get("title") or source_title)
        source_link = self._source_link(source_label)
        source_link_html = f'<p class="source"><a href="{html.escape(source_link)}">{html.escape(source_link)}</a></p>' if source_link else ""
        scene_cards = []
        for scene in scenes:
            rel_image = scene.image_path.relative_to(html_path.parent).as_posix() if scene.image_path else ""
            script_html = "".join(f"<p>{html.escape(paragraph)}</p>" for paragraph in self._note_paragraphs(scene.script))
            scene_cards.append(
                f"""
                <section class="scene">
                  <h2>{html.escape(scene.heading)} ({html.escape(scene.timecode)})</h2>
                  <img src="{html.escape(rel_image)}" alt="{html.escape(scene.heading)}">
                  <div class="script">{script_html}</div>
                </section>
                """
            )
        document = f"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      color-scheme: light;
      --ink: #17202a;
      --muted: #5f6b7a;
      --line: #dbe2ea;
      --paper: #fbfcfe;
      --accent: #1677ff;
    }}
    body {{
      margin: 0;
      font-family: "Segoe UI", "Malgun Gothic", sans-serif;
      background: var(--paper);
      color: var(--ink);
      line-height: 1.65;
    }}
    main {{
      width: min(960px, calc(100vw - 40px));
      margin: 40px auto 72px;
    }}
    header {{
      border-bottom: 1px solid var(--line);
      padding-bottom: 24px;
      margin-bottom: 28px;
    }}
    h1 {{
      margin: 0 0 12px;
      font-size: clamp(28px, 5vw, 44px);
      letter-spacing: 0;
      line-height: 1.15;
    }}
    .source {{
      color: var(--muted);
      font-size: 14px;
      overflow-wrap: anywhere;
    }}
    .credit {{
      display: inline-flex;
      margin-top: 10px;
      padding: 3px 9px;
      border-radius: 6px;
      background: #eaf2ff;
      color: #2563eb;
      font-size: 13px;
      font-weight: 700;
    }}
    .scene {{
      padding: 32px 0 38px;
      border-top: 1px solid var(--line);
      break-inside: avoid;
    }}
    h2 {{
      margin: 0 0 16px;
      font-size: 24px;
      letter-spacing: 0;
    }}
    img {{
      width: 100%;
      max-width: 760px;
      max-height: 520px;
      object-fit: contain;
      display: block;
      margin: 0 auto;
      background: #111827;
      border-radius: 8px;
      border: 1px solid var(--line);
    }}
    .script {{
      margin-top: 18px;
    }}
    p {{
      margin: 0 0 12px;
    }}
    footer {{
      border-top: 1px solid var(--line);
      color: var(--muted);
      font-size: 13px;
      padding-top: 18px;
      margin-top: 28px;
    }}
  </style>
</head>
<body>
  <main>
    <header>
      <h1>{html.escape(title)}</h1>
      {source_link_html}
    </header>
    {''.join(scene_cards)}
    <footer>영상·음성 요약 노트 생성기 · developed by yeohj0710</footer>
  </main>
</body>
</html>
"""
        html_path.write_text(document, encoding="utf-8")
        return html_path

    def _register_pdf_fonts(self) -> tuple[str, str]:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        windows = Path(os.getenv("WINDIR", r"C:\Windows")) / "Fonts"
        regular_candidates = [
            windows / "malgun.ttf",
            windows / "NanumGothic.ttf",
            Path(r"C:\Windows\Fonts\malgun.ttf"),
        ]
        bold_candidates = [
            windows / "malgunbd.ttf",
            windows / "NanumGothicBold.ttf",
            Path(r"C:\Windows\Fonts\malgunbd.ttf"),
        ]

        regular_path = next((path for path in regular_candidates if path.exists()), None)
        bold_path = next((path for path in bold_candidates if path.exists()), regular_path)
        if not regular_path:
            return "Helvetica", "Helvetica-Bold"

        if "ClipNoteKorean" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ClipNoteKorean", str(regular_path)))
        if bold_path and "ClipNoteKorean-Bold" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ClipNoteKorean-Bold", str(bold_path)))
        return "ClipNoteKorean", "ClipNoteKorean-Bold" if bold_path else "ClipNoteKorean"

    def _pdf_image(self, image_path: Path, max_width: float, max_height: float):
        from PIL import Image as PILImage
        from reportlab.platypus import Image

        with PILImage.open(image_path) as image:
            width, height = image.size
        if width <= 0 or height <= 0:
            return None
        scale = min(max_width / width, max_height / height, 1.0)
        return Image(str(image_path), width=width * scale, height=height * scale)

    def _pdf_paragraph_text(self, value: object) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        return html.escape(text).replace("\n", "<br/>")

    def _render_pdf(
        self,
        job_dir: Path,
        source_title: str,
        source_label: str,
        duration: float,
        chunks: list[TranscriptChunk],
        scenes: list[Scene],
        analysis: dict[str, object],
    ) -> Path:
        self.progress("PDF 생성 중", 0.94, "장면 이미지와 대본을 PDF로 정리하고 있습니다.")

        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer

        pdf_path = job_dir / USER_PDF_NAME
        regular_font, bold_font = self._register_pdf_fonts()
        title = str(analysis.get("title") or source_title)
        source_link = self._source_link(source_label)

        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=17 * mm,
            bottomMargin=17 * mm,
            title=title,
            author="영상·음성 요약 노트 생성기",
            subject="developed by yeohj0710",
        )
        width, _height = A4
        content_width = width - doc.leftMargin - doc.rightMargin

        styles = {
            "title": ParagraphStyle(
                "ClipNoteTitle",
                fontName=bold_font,
                fontSize=22,
                leading=30,
                textColor=colors.HexColor("#142033"),
                spaceAfter=10,
            ),
            "h2": ParagraphStyle(
                "ClipNoteH2",
                fontName=bold_font,
                fontSize=15.5,
                leading=22,
                textColor=colors.HexColor("#142033"),
                spaceBefore=4,
                spaceAfter=9,
            ),
            "h3": ParagraphStyle(
                "ClipNoteH3",
                fontName=bold_font,
                fontSize=12,
                leading=18,
                textColor=colors.HexColor("#1677ff"),
                spaceBefore=8,
                spaceAfter=5,
            ),
            "body": ParagraphStyle(
                "ClipNoteBody",
                fontName=regular_font,
                fontSize=10.2,
                leading=16,
                textColor=colors.HexColor("#17202a"),
                spaceAfter=7,
            ),
            "muted": ParagraphStyle(
                "ClipNoteMuted",
                fontName=regular_font,
                fontSize=9,
                leading=14,
                textColor=colors.HexColor("#5f6b7a"),
                spaceAfter=6,
            ),
            "link": ParagraphStyle(
                "ClipNoteLink",
                fontName=regular_font,
                fontSize=9,
                leading=14,
                textColor=colors.HexColor("#2563eb"),
                spaceAfter=9,
            ),
        }

        story = [
            Paragraph(self._pdf_paragraph_text(title), styles["title"]),
        ]
        if source_link:
            story.append(Paragraph(self._pdf_paragraph_text(source_link), styles["link"]))
        story.append(Spacer(1, 4 * mm))

        for index, scene in enumerate(scenes):
            if index > 0:
                story.append(PageBreak())
            block = [Paragraph(f"{self._pdf_paragraph_text(scene.heading)} ({scene.timecode})", styles["h2"])]
            if scene.image_path and scene.image_path.exists():
                image = self._pdf_image(scene.image_path, content_width * 0.88, 88 * mm)
                if image:
                    image.hAlign = "CENTER"
                    block.append(image)
                    block.append(Spacer(1, 4 * mm))
            story.append(KeepTogether(block))
            for paragraph in self._note_paragraphs(scene.script):
                story.append(Paragraph(self._pdf_paragraph_text(paragraph), styles["body"]))

        def draw_footer(canvas, document):
            canvas.saveState()
            canvas.setFont(regular_font, 8)
            canvas.setFillColor(colors.HexColor("#7a8797"))
            canvas.drawRightString(width - doc.rightMargin, 9 * mm, f"영상·음성 요약 노트 생성기 · developed by yeohj0710 · {document.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
        return pdf_path

    def _render_docx(
        self,
        job_dir: Path,
        source_title: str,
        source_label: str,
        scenes: list[Scene],
        analysis: dict[str, object],
    ) -> Path:
        self.progress("DOCX 생성 중", 0.97, "노션에 붙여넣기 쉬운 Word 문서를 만들고 있습니다.")

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        from PIL import Image as PILImage

        docx_path = job_dir / USER_DOCX_NAME
        title = str(analysis.get("title") or source_title)
        source_link = self._source_link(source_label)

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "맑은 고딕"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        normal.font.size = Pt(10.5)

        title_style = styles["Title"]
        title_style.font.name = "맑은 고딕"
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        title_style.font.size = Pt(22)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(20, 32, 51)

        heading_style = styles["Heading 1"]
        heading_style.font.name = "맑은 고딕"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        heading_style.font.size = Pt(15)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(20, 32, 51)

        document.add_paragraph(f"# {title}")
        if source_link:
            paragraph = document.add_paragraph(source_link)
            run = paragraph.runs[0]
            run.font.color.rgb = RGBColor(37, 99, 235)
            run.font.size = Pt(9)

        max_width = 5.4
        max_height = 3.7
        for index, scene in enumerate(scenes):
            if index > 0:
                document.add_page_break()
            document.add_paragraph(f"## {scene.heading} ({scene.timecode})")
            if scene.image_path and scene.image_path.exists():
                with PILImage.open(scene.image_path) as image:
                    width_px, height_px = image.size
                if width_px > 0 and height_px > 0:
                    width_inches = max_width
                    height_inches = width_inches * height_px / width_px
                    if height_inches > max_height:
                        height_inches = max_height
                        width_inches = height_inches * width_px / height_px
                    picture_paragraph = document.add_paragraph()
                    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    picture_paragraph.add_run().add_picture(str(scene.image_path), width=Inches(width_inches))
            for paragraph_text in self._note_paragraphs(scene.script):
                paragraph = document.add_paragraph(paragraph_text)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.15

        document.core_properties.author = "yeohj0710"
        document.save(docx_path)
        return docx_path

    def _write_metadata(
        self,
        job_dir: Path,
        source_title: str,
        source_label: str,
        duration: float,
        scenes: list[Scene],
        analysis: dict[str, object],
    ) -> None:
        payload = {
            "source_title": source_title,
            "source": source_label,
            "duration_seconds": duration,
            "transcription_model": self.settings.transcription_model,
            "text_model": self.settings.text_model,
            "analysis": analysis,
            "scenes": [
                {
                    "index": scene.index,
                    "seconds": scene.seconds,
                    "timecode": scene.timecode,
                    "heading": scene.heading,
                    "summary": scene.summary,
                    "quote": scene.quote,
                    "why": scene.why,
                    "script": scene.script,
                    "image": str(scene.image_path.relative_to(job_dir)) if scene.image_path else "",
                }
                for scene in scenes
            ],
        }
        (job_dir / "metadata.json").write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

